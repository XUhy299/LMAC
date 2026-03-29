"""
核心功能模块 - 文档加载、文本分块、向量存储（ES混合检索）、重排序、LLM封装
"""
import os
import re
import json
import time
import pickle
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime

import numpy as np
import requests
import fitz  # PyMuPDF

# Elasticsearch
from elasticsearch import Elasticsearch
from elasticsearch.helpers import bulk

# LangChain相关
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document as LangChainDocument

# Transformers - CrossEncoder重排序
from sentence_transformers import CrossEncoder

# LlamaIndex - 高级文档处理
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SentenceSplitter

import agentic_rag.config as config


# ==================== 数据类定义 ====================

@dataclass
class SearchResult:
    """搜索结果数据类"""
    chunk_id: str
    content: str
    score: float
    metadata: Dict[str, Any] = field(default_factory=dict)
    bm25_score: float = 0.0  # BM25 分数
    vector_score: float = 0.0  # 向量相似度分数


@dataclass
class RAGResult:
    """RAG结果数据类"""
    original_query: str
    sub_questions: List[str] = field(default_factory=list)
    sub_answers: Dict[str, str] = field(default_factory=dict)
    final_answer: str = ""
    sources: List[SearchResult] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


# ==================== 文档处理模块 ====================

class DocumentLoader:
    """
    文档加载器 - 支持PDF等格式
    """

    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.md', '.docx', '.doc']

    def load_pdf(self, file_path: str) -> List[LangChainDocument]:
        """加载PDF文件"""
        try:
            loader = PyMuPDFLoader(file_path)
            documents = loader.load()

            # 清洗文本
            for doc in documents:
                doc.page_content = self._clean_text(doc.page_content)

            return documents
        except Exception as e:
            print(f"加载PDF失败: {e}")
            return []

    def load_txt(self, file_path: str) -> List[LangChainDocument]:
        """加载文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            documents = [LangChainDocument(
                page_content=content,
                metadata={"source": file_path, "type": "txt"}
            )]
            return documents
        except Exception as e:
            print(f"加载文本文件失败: {e}")
            return []

    def load_docx(self, file_path: str) -> List[LangChainDocument]:
        """加载Word文档"""
        try:
            from docx import Document
            doc = Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            documents = [LangChainDocument(
                page_content=self._clean_text(content),
                metadata={"source": file_path, "type": "docx"}
            )]
            return documents
        except ImportError:
            print("未安装 python-docx，无法读取 .docx 文件")
            return []
        except Exception as e:
            print(f"加载Word文档失败: {e}")
            return []

    def load_directory(self, directory: str, file_types: List[str] = None) -> List[LangChainDocument]:
        """加载目录下的所有文档"""
        all_documents = []

        if file_types is None:
            file_types = config.KNOWLEDGE_FILE_TYPES

        for root, dirs, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)
                ext = os.path.splitext(file)[1].lower()

                if ext not in file_types:
                    continue

                if ext == '.pdf':
                    docs = self.load_pdf(file_path)
                elif ext in ['.txt', '.md']:
                    docs = self.load_txt(file_path)
                elif ext in ['.docx', '.doc']:
                    docs = self.load_docx(file_path)
                else:
                    continue

                # 添加文件元信息
                for doc in docs:
                    doc.metadata["filename"] = file
                    doc.metadata["file_path"] = file_path
                    doc.metadata["file_type"] = ext

                all_documents.extend(docs)

        return all_documents

    def _clean_text(self, text: str) -> str:
        """中英文混合文本清洗"""
        # 处理 PDF 软换行
        text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)

        # 统一标点
        text = re.sub(r'[^\w\s\u4e00-\u9fff\u3000-\u303f.,;:!?()\[\]{}"\'`~@#$%^&*+=|\\/<>-]', ' ', text)

        # 引用过滤
        text = re.sub(r'\[\d+(?:[,\-\s]+\d+)*\]', '', text)
        text = re.sub(r'\(\d+(?:[,\-\s]+\d+)*\)', '', text)

        # 中英文排版优化
        text = re.sub(r'([\u4e00-\u9fff])([a-zA-Z0-9])', r'\1 \2', text)
        text = re.sub(r'([a-zA-Z0-9])([\u4e00-\u9fff])', r'\1 \2', text)

        # 清理空白字符
        text = re.sub(r'[\r\n\t]+', ' ', text)
        text = re.sub(r'\s{2,}', ' ', text)

        return text.strip()


class TextSplitter:
    """
    文本分块器
    """

    def __init__(
        self,
        chunk_size: int = config.CHUNK_SIZE,
        chunk_overlap: int = config.CHUNK_OVERLAP
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # LangChain递归字符分块器
        self.langchain_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
            length_function=len
        )

        # LlamaIndex句子分块器
        self.llama_splitter = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )

    def split_documents(self, documents: List[LangChainDocument]) -> List[LangChainDocument]:
        """分块文档"""
        return self.langchain_splitter.split_documents(documents)

    def split_text(self, text: str) -> List[str]:
        """分块文本"""
        return self.langchain_splitter.split_text(text)


# ==================== Elasticsearch 向量存储模块 ====================

class ElasticsearchVectorStore:
    """
    Elasticsearch 向量存储 - 支持混合检索（BM25 + 向量相似度）
    """

    def __init__(
        self,
        host: str = config.ES_HOST,
        port: int = config.ES_PORT,
        user: str = config.ES_USER,
        password: str = config.ES_PASSWORD,
        index_name: str = config.ES_INDEX_NAME,
        embedding_model: str = config.EMBEDDING_MODEL
    ):
        self.host = host
        self.port = port
        self.user = user
        self.password = password
        self.index_name = index_name
        self.vector_dim = config.ES_VECTOR_DIM

        # 初始化 embedding 模型
        self.embeddings = HuggingFaceEmbeddings(
            model_name=embedding_model,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )

        # 初始化 ES 客户端
        self.es = self._create_client()

        # 确保索引存在
        self._ensure_index()

    def _create_client(self) -> Elasticsearch:
        """创建 ES 客户端"""
        es_config = {
            "hosts": [f"http://{self.host}:{self.port}"],
            "retry_on_timeout": True,
            "max_retries": 3,
            "timeout": 30
        }

        if self.user and self.password:
            es_config["basic_auth"] = (self.user, self.password)

        return Elasticsearch(**es_config)

    def _ensure_index(self):
        """确保索引存在，不存在则创建"""
        if not self.es.indices.exists(index=self.index_name):
            self._create_index()

    def _create_index(self):
        """创建 ES 索引，配置混合检索"""
        mapping = {
            "mappings": {
                "properties": {
                    "content": {
                        "type": "text",
                        "analyzer": "standard",
                        "fields": {
                            "keyword": {
                                "type": "keyword"
                            }
                        }
                    },
                    "embedding": {
                        "type": "dense_vector",
                        "dims": self.vector_dim,
                        "index": True,
                        "similarity": "cosine"
                    },
                    "metadata": {
                        "type": "object",
                        "properties": {
                            "filename": {"type": "keyword"},
                            "file_path": {"type": "keyword"},
                            "file_type": {"type": "keyword"},
                            "chunk_index": {"type": "integer"},
                            "knowledge_base_id": {"type": "keyword"}
                        }
                    },
                    "chunk_id": {"type": "keyword"},
                    "created_at": {"type": "date"}
                }
            },
            "settings": {
                "number_of_shards": 1,
                "number_of_replicas": 0,
                "index": {
                    "similarity": {
                        "default": {
                            "type": "BM25"
                        }
                    }
                }
            }
        }

        self.es.indices.create(index=self.index_name, body=mapping)
        print(f"ES 索引创建成功: {self.index_name}")

    def add_documents(
        self,
        documents: List[LangChainDocument],
        knowledge_base_id: str = "default"
    ) -> List[str]:
        """
        添加文档到 ES
        返回添加的文档 ID 列表
        """
        actions = []
        doc_ids = []

        for i, doc in enumerate(documents):
            # 生成 embedding
            embedding = self.embeddings.embed_query(doc.page_content)

            # 生成唯一 chunk_id
            chunk_id = hashlib.md5(
                f"{doc.metadata.get('file_path', '')}_{i}_{doc.page_content[:100]}".encode()
            ).hexdigest()

            # 准备元数据
            metadata = doc.metadata.copy()
            metadata["chunk_index"] = i
            metadata["knowledge_base_id"] = knowledge_base_id

            action = {
                "_index": self.index_name,
                "_id": chunk_id,
                "_source": {
                    "content": doc.page_content,
                    "embedding": embedding,
                    "metadata": metadata,
                    "chunk_id": chunk_id,
                    "created_at": datetime.now().isoformat()
                }
            }

            actions.append(action)
            doc_ids.append(chunk_id)

        # 批量写入
        if actions:
            try:
                success, errors = bulk(self.es, actions, refresh=True, raise_on_error=False)
                failed_count = len(errors) if errors else 0
                print(f"ES 批量写入: 成功 {success}, 失败 {failed_count}")
            except Exception as e:
                print(f"ES 批量写入异常: {e}")
                return []

        return doc_ids

    def hybrid_search(
        self,
        query: str,
        top_k: int = config.ES_TOP_K,
        knowledge_base_id: str = None,
        min_score: float = config.ES_MIN_SCORE
    ) -> List[SearchResult]:
        """
        混合检索：BM25 + 向量相似度
        """
        # 生成查询向量
        query_embedding = self.embeddings.embed_query(query)

        # 构建混合检索查询
        search_body = {
            "size": top_k,
            "query": {
                "bool": {
                    "should": [
                        # BM25 文本检索
                        {
                            "multi_match": {
                                "query": query,
                                "fields": ["content^2", "content.keyword"],
                                "type": "best_fields",
                                "boost": config.ES_BM25_WEIGHT
                            }
                        },
                        # 向量相似度检索
                        {
                            "script_score": {
                                "query": {"match_all": {}},
                                "script": {
                                    "source": f"cosineSimilarity(params.query_vector, 'embedding') + 1.0",
                                    "params": {"query_vector": query_embedding}
                                },
                                "boost": config.ES_VECTOR_WEIGHT
                            }
                        }
                    ]
                }
            },
            "_source": ["content", "metadata", "chunk_id"]
        }

        # 添加知识库过滤
        if knowledge_base_id:
            search_body["query"]["bool"]["filter"] = [
                {"term": {"metadata.knowledge_base_id": knowledge_base_id}}
            ]

        # 执行检索
        response = self.es.search(index=self.index_name, body=search_body)

        # 解析结果
        results = []
        for hit in response["hits"]["hits"]:
            score = hit["_score"]
            if score < min_score:
                continue

            source = hit["_source"]
            # 混合检索分数：cosineSimilarity + 1.0 后需要归一化到 0-1
            # 但 BM25 分数无固定范围，这里只归一化向量部分
            vector_score = (score - 1.0) / 2.0 if score > 1.0 else score / 2.0
            results.append(SearchResult(
                chunk_id=source.get("chunk_id", hit["_id"]),
                content=source["content"],
                score=max(0.0, min(1.0, vector_score)),  # 限制在 0-1 范围
                metadata=source.get("metadata", {}),
                bm25_score=0.0,  # ES 不单独返回 BM25 分数
                vector_score=max(0.0, min(1.0, vector_score))
            ))

        return results

    def vector_search(
        self,
        query: str,
        top_k: int = config.TOP_K,
        knowledge_base_id: str = None
    ) -> List[SearchResult]:
        """
        纯向量检索
        """
        query_embedding = self.embeddings.embed_query(query)

        search_body = {
            "size": top_k,
            "query": {
                "script_score": {
                    "query": {"match_all": {}},
                    "script": {
                        "source": "cosineSimilarity(params.query_vector, 'embedding') + 1.0",
                        "params": {"query_vector": query_embedding}
                    }
                }
            },
            "_source": ["content", "metadata", "chunk_id"]
        }

        if knowledge_base_id:
            search_body["query"]["script_score"]["query"] = {
                "term": {"metadata.knowledge_base_id": knowledge_base_id}
            }

        response = self.es.search(index=self.index_name, body=search_body)

        results = []
        for hit in response["hits"]["hits"]:
            source = hit["_source"]
            results.append(SearchResult(
                chunk_id=source.get("chunk_id", hit["_id"]),
                content=source["content"],
                score=(hit["_score"] - 1.0) / 2.0,  # 转换为 0-1 范围
                metadata=source.get("metadata", {}),
                bm25_score=0.0,
                vector_score=(hit["_score"] - 1.0) / 2.0
            ))

        return results

    def delete_by_knowledge_base(self, knowledge_base_id: str):
        """删除指定知识库的所有文档"""
        query = {
            "query": {
                "term": {"metadata.knowledge_base_id": knowledge_base_id}
            }
        }
        self.es.delete_by_query(index=self.index_name, body=query, refresh=True)

    def get_stats(self) -> Dict[str, Any]:
        """获取索引统计信息"""
        stats = self.es.indices.stats(index=self.index_name)
        return {
            "total_docs": stats["indices"][self.index_name]["total"]["docs"]["count"],
            "index_name": self.index_name
        }


# ==================== 检索器模块 ====================

class CustomRetriever:
    """
    自定义检索器 - 结合 ES 混合检索和重排序
    """

    def __init__(
        self,
        vectorstore: ElasticsearchVectorStore,
        reranker_model: Optional[CrossEncoder] = None,
        top_k: int = config.TOP_K,
        rerank_top_k: int = config.RERANK_CANDIDATES
    ):
        self.vectorstore = vectorstore
        self.reranker_model = reranker_model
        self.top_k = top_k
        self.rerank_top_k = rerank_top_k

    def get_relevant_documents(
        self,
        query: str,
        use_reranker: bool = True,
        knowledge_base_id: str = None
    ) -> List[SearchResult]:
        """获取相关文档"""
        # 初步检索 - 获取更多候选文档
        initial_k = self.rerank_top_k if use_reranker and self.reranker_model else self.top_k

        results = self.vectorstore.hybrid_search(
            query=query,
            top_k=initial_k,
            knowledge_base_id=knowledge_base_id
        )

        if not use_reranker or not self.reranker_model:
            return results[:self.top_k]

        # 使用 CrossEncoder 重排序
        return self._rerank_documents(query, results)

    def _rerank_documents(
        self,
        query: str,
        docs: List[SearchResult]
    ) -> List[SearchResult]:
        """使用 CrossEncoder 重排序"""
        if not docs:
            return []

        # 准备重排序输入
        doc_contents = [doc.content for doc in docs]
        pairs = [[query, doc] for doc in doc_contents]

        try:
            scores = self.reranker_model.predict(pairs)

            # 归一化分数到 0-1 范围
            scores = np.array(scores)
            scores_shifted = scores - np.max(scores)
            normalized_scores = np.exp(scores_shifted) / np.exp(scores_shifted).sum()

            # 组合结果
            for i, doc in enumerate(docs):
                doc.score = float(normalized_scores[i])

            # 按分数排序
            docs.sort(key=lambda x: x.score, reverse=True)

            return docs[:self.top_k]

        except Exception as e:
            print(f"重排序失败: {e}")
            return docs[:self.top_k]


# ==================== 重排序模块 ====================

class Reranker:
    """
    文档重排序器 - 使用 Transformers CrossEncoder
    """

    def __init__(self, model_name: str = config.RERANKER_MODEL):
        self.model_name = model_name
        self.model = None
        self._load_model()

    def _load_model(self):
        """加载 CrossEncoder 模型"""
        try:
            self.model = CrossEncoder(self.model_name)
            print(f"重排序模型加载成功: {self.model_name}")
        except Exception as e:
            print(f"加载重排序模型失败: {e}")
            self.model = None

    def rerank(
        self,
        query: str,
        documents: List[str],
        top_k: int = config.TOP_K
    ) -> List[Tuple[int, float]]:
        """
        重排序文档
        返回: [(文档索引, 归一化分数), ...]
        """
        if not documents or self.model is None:
            return [(i, 1.0 / len(documents)) for i in range(len(documents))]

        # 构建查询-文档对
        pairs = [[query, doc] for doc in documents]

        # 预测分数
        scores = self.model.predict(pairs)

        # 归一化
        scores = np.array(scores)
        scores_shifted = scores - np.max(scores)
        normalized_scores = np.exp(scores_shifted) / np.exp(scores_shifted).sum()

        # 组合索引和分数
        indexed_scores = list(enumerate(normalized_scores))

        # 按分数降序排序
        indexed_scores.sort(key=lambda x: x[1], reverse=True)

        return indexed_scores[:top_k]


# ==================== LLM封装模块 ====================

class LLMWrapper:
    """
    LLM 封装器 - 支持 Ollama
    """

    def __init__(
        self,
        base_url: str = config.OLLAMA_BASE_URL,
        default_model: str = config.ANSWER_MODEL
    ):
        self.base_url = base_url
        self.default_model = default_model

        # Ollama API 端点
        self.embedding_url = f"{base_url}/api/embeddings"
        self.generate_url = f"{base_url}/api/generate"

    def _generate_with_retry(
        self,
        payload: Dict[str, Any],
        max_retries: int = 3
    ) -> Dict[str, Any]:
        """带重试的生成方法"""
        for attempt in range(max_retries):
            try:
                response = requests.post(
                    self.generate_url,
                    json=payload,
                    timeout=120
                )
                response.raise_for_status()
                return response.json()
            except Exception as e:
                if attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"[重试] 生成失败，{wait_time}秒后重试 ({attempt + 1}/{max_retries}): {e}")
                    time.sleep(wait_time)
                else:
                    raise e
        return {}

    def generate(
        self,
        prompt: str,
        model: str = None,
        system_prompt: str = "",
        temperature: float = 0.7,
        max_tokens: int = 4096,
        stream: bool = False
    ) -> Dict[str, Any]:
        """
        生成文本
        返回: {"response": str, "duration": float, "model": str}
        """
        model = model or self.default_model

        if config.MOCK_LLM:
            return self._mock_generate(prompt, model)

        full_prompt = f"{system_prompt}\n\n{prompt}" if system_prompt else prompt

        payload = {
            "model": model,
            "prompt": full_prompt,
            "stream": stream,
            "temperature": temperature,
            "top_p": 0.9,
            "top_k": 40,
            "num_predict": max_tokens
        }

        start_time = time.time()

        try:
            if stream:
                response_text = ""
                with requests.post(
                    self.generate_url,
                    json=payload,
                    timeout=120,
                    stream=True
                ) as response:
                    response.raise_for_status()
                    for line in response.iter_lines():
                        if line:
                            data = json.loads(line.decode('utf-8'))
                            token = data.get("response", "")
                            response_text += token

                            if data.get("done", False):
                                break

                duration = time.time() - start_time
                return {
                    "response": response_text,
                    "duration": duration,
                    "model": model
                }
            else:
                result = self._generate_with_retry(payload)

                duration = time.time() - start_time
                return {
                    "response": result.get("response", "") if result else "",
                    "duration": duration,
                    "model": model
                }

        except Exception as e:
            duration = time.time() - start_time
            return {
                "response": f"生成失败: {str(e)}",
                "duration": duration,
                "model": model
            }

    def chat(
        self,
        messages: List[Dict[str, str]],
        model: str = None,
        temperature: float = 0.7
    ) -> Dict[str, Any]:
        """
        对话接口
        messages: [{"role": "user/assistant/system", "content": "..."}]
        """
        model = model or self.default_model

        system_prompt = ""
        user_prompt = ""

        for msg in messages:
            if msg["role"] == "system":
                system_prompt = msg["content"]
            elif msg["role"] == "user":
                user_prompt = msg["content"]

        return self.generate(
            prompt=user_prompt,
            model=model,
            system_prompt=system_prompt,
            temperature=temperature
        )

    def _mock_generate(self, prompt: str, model: str) -> Dict[str, Any]:
        """Mock LLM - 测试时不需要真实 API"""
        import re

        # 提取关键词
        words = re.findall(r'[\u4e00-\u9fff]{2,}|[a-zA-Z]{3,}', prompt)
        keywords = ' '.join(words[:5]) if words else '该主题'

        if any(kw in prompt for kw in ['分解', 'sub-question', '子问题', 'decompose']):
            response = json.dumps([
                "这个问题的第一个方面是什么？",
                "相关的背景知识是什么？",
                "有哪些实际应用？"
            ], ensure_ascii=False)
        elif any(kw in prompt for kw in ['综合', 'synthesize', '最终', '总结', 'synthesis']):
            response = f"综合以上分析，{keywords} 的核心要点包括：首先，它具有重要的理论意义；其次，在实践中有广泛应用；最后，未来发展前景广阔。"
        elif any(kw in prompt for kw in ['事实核查', 'fact check', '验证', 'verify']):
            response = prompt
        elif any(kw in prompt for kw in ['回答', 'answer', '请回答', '问题']):
            response = f"根据文档内容，{keywords} 是指一种重要的技术或概念，它在相关领域发挥着关键作用。具体来说，它通过特定机制实现了预期功能，为用户提供了有价值的服务。"
        else:
            response = f"[Mock回答] 关于 '{keywords}' 的问题：这是一个模拟响应，用于测试目的。实际部署时将调用真实的LLM API。"

        return {
            "response": response,
            "duration": 0.01,
            "model": f"mock-{model}"
        }


# ==================== 工具函数 ====================

def compute_cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """计算余弦相似度"""
    vec1 = np.array(vec1)
    vec2 = np.array(vec2)

    dot_product = np.dot(vec1, vec2)
    norm1 = np.linalg.norm(vec1)
    norm2 = np.linalg.norm(vec2)

    if norm1 == 0 or norm2 == 0:
        return 0.0

    return float(dot_product / (norm1 * norm2))


def get_text_hash(text: str) -> str:
    """获取文本哈希"""
    return hashlib.md5(text.encode()).hexdigest()


def save_cache(data: Any, cache_file: str):
    """保存缓存"""
    os.makedirs(os.path.dirname(cache_file), exist_ok=True)
    with open(cache_file, 'wb') as f:
        pickle.dump(data, f)


def load_cache(cache_file: str) -> Optional[Any]:
    """加载缓存"""
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'rb') as f:
                return pickle.load(f)
        except Exception as e:
            print(f"加载缓存失败: {e}")
    return None


# ==================== 主程序入口 ====================

def initialize_system(knowledge_base_path: str = None):
    """
    初始化系统 - 加载文档、创建向量存储
    返回: (vectorstore, documents, text_chunks, reranker)
    """
    print("=" * 50)
    print("初始化 Agentic-RAG 系统")
    print("=" * 50)

    # 创建知识库目录
    kb_path = knowledge_base_path or config.KNOWLEDGE_BASE_DIR
    os.makedirs(kb_path, exist_ok=True)

    # 初始化 ES 向量存储
    print("\n[1/4] 初始化 Elasticsearch 向量存储...")
    try:
        vectorstore = ElasticsearchVectorStore()
        stats = vectorstore.get_stats()
        print(f"ES 索引已连接，当前文档数: {stats['total_docs']}")
    except Exception as e:
        print(f"ES 连接失败: {e}")
        raise

    # 加载文档
    print(f"\n[2/4] 加载知识库文档 from {kb_path}...")
    loader = DocumentLoader()
    documents = loader.load_directory(kb_path)
    print(f"加载了 {len(documents)} 个文档")

    if not documents:
        print("警告: 未找到文档，使用示例文本")
        example_text = """
        RAG (Retrieval-Augmented Generation) 检索增强生成，是一种结合检索系统和生成模型的技术。
        它通过从外部知识库中检索相关信息，然后将这些信息作为上下文提供给生成模型，从而生成更准确、更相关的回答。

        Agentic-RAG 是RAG的进一步发展，引入了智能代理(Agent)的概念。
        它的核心思想是让RAG系统具有自主决策能力，能够：
        1. 分析用户问题，决定是否需要检索
        2. 将复杂问题分解为多个子问题
        3. 制定检索策略，选择合适的检索方式
        4. 评估检索结果，必要时进行多轮检索
        5. 综合多个答案生成最终回复

        Agentic-RAG的优势：
        - 能够处理更复杂的问题
        - 检索结果更精准
        - 回答更加全面和准确
        - 具有自我纠错能力
        """
        documents = [LangChainDocument(
            page_content=example_text,
            metadata={"source": "example", "type": "example"}
        )]

    # 文本分块
    print("\n[3/4] 文本分块...")
    splitter = TextSplitter()
    text_chunks = splitter.split_documents(documents)
    print(f"生成了 {len(text_chunks)} 个文本块")

    # 添加到 ES（如果知识库路径有文档）
    if knowledge_base_path and text_chunks:
        print("\n[4/4] 索引文档到 ES...")
        vectorstore.add_documents(text_chunks, knowledge_base_id="default")
        print(f"文档索引完成")

    # 初始化重排序模型
    print("\n加载重排序模型...")
    try:
        reranker = Reranker()
        print("重排序模型加载成功")
    except Exception as e:
        print(f"重排序模型加载失败: {e}")
        reranker = None

    print("\n" + "=" * 50)
    print("系统初始化完成！")
    print("=" * 50)

    return vectorstore, documents, text_chunks, reranker


def get_summary_llm():
    """
    获取用于 ConversationSummaryBufferMemory 的 LLM
    """
    return LLMWrapper(
        base_url=config.OLLAMA_BASE_URL,
        default_model=config.SUMMARY_LLM_MODEL
    )
